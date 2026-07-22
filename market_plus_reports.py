from __future__ import annotations

import math
import tempfile
from datetime import date, datetime
from html import escape as xml_escape
from pathlib import Path
from typing import Any, Iterable

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


NAVY = "17365D"
PURPLE = "6C31D9"
BLUE = "1F9AD6"
PALE_BLUE = "EAF4FA"
PALE_GRAY = "F5F7FA"
GREEN = "2E9B4F"
RED = "C0392B"
WHITE = "FFFFFF"
GRID = Side(style="thin", color="D9E2EA")


def _num(value: Any) -> float | None:
    if value is None or isinstance(value, bool):
        return None
    try:
        result = float(value)
    except (TypeError, ValueError):
        return None
    return result if math.isfinite(result) else None


def _money(value: Any) -> str:
    number = _num(value)
    return "—" if number is None else f"{number:,.0f}".replace(",", " ")


def _signed_money(value: Any, sign: str) -> str:
    number = max(0.0, float(_num(value) or 0))
    return _money(0) if number <= 0 else f"{sign} {_money(number)}"


def _qty(value: Any) -> int:
    number = _num(value)
    return max(0, int(number or 0))


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _identifier(value: Any) -> str:
    if value is None or isinstance(value, bool):
        return ""
    if isinstance(value, (int, float)):
        number = float(value)
        if math.isfinite(number) and number.is_integer():
            return str(int(number))
    text = _clean(value)
    if text.endswith(".0"):
        try:
            number = float(text)
            if number.is_integer():
                return str(int(number))
        except ValueError:
            pass
    return text


def _date_label(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return _clean(value) or "—"


def _resolve_output(output: str | Path | None, prefix: str, suffix: str) -> Path:
    if output is None:
        handle = tempfile.NamedTemporaryFile(prefix=prefix, suffix=suffix, delete=False)
        handle.close()
        return Path(handle.name)
    path = Path(output)
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def prepare_compensation_rows(
    stock_rows: Iterable[dict[str, Any]],
    *,
    kind: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Build claim rows only from documented Uzum catalog values.

    The attached Market Plus claim explains compensation as current value less
    marketplace commission.  We calculate that value only when both the sale
    price and exact SKU commission are present.  A commission range is never
    guessed.
    """

    normalized_kind = str(kind).strip().lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    quantity_field = "missing" if normalized_kind == "loss" else "defected"
    rows: list[dict[str, Any]] = []
    missing_compensation = 0
    total_qty = 0
    total_value = 0.0

    for source in stock_rows:
        qty = _qty(source.get(quantity_field))
        if qty <= 0:
            continue
        price = _num(source.get("price"))
        commission = _num(source.get("commission"))
        compensation: float | None = None
        if (
            price is not None
            and price > 0
            and commission is not None
            and 0 <= commission <= 100
        ):
            compensation = max(0.0, price * (1.0 - commission / 100.0))
        else:
            missing_compensation += 1

        total = compensation * qty if compensation is not None else None
        total_qty += qty
        if total is not None:
            total_value += total
        rows.append(
            {
                "product_id": _identifier(source.get("product_id")),
                "title": _clean(
                    source.get("product_title")
                    or source.get("sku_full_title")
                    or source.get("sku_title")
                    or "Без названия"
                ),
                "sku": _clean(
                    source.get("seller_item_code")
                    or source.get("sku_full_title")
                    or source.get("sku_title")
                    or source.get("sku_id")
                ),
                "sku_id": _identifier(source.get("sku_id")),
                "barcode": _identifier(source.get("barcode")),
                "price": price,
                "commission_percent": commission,
                "compensation": compensation,
                "quantity": qty,
                "total": total,
            }
        )

    rows.sort(key=lambda item: (-int(item["quantity"]), str(item["title"]).lower()))
    return rows, {
        "kind": normalized_kind,
        "rows": len(rows),
        "quantity": total_qty,
        "known_total": total_value,
        "missing_compensation": missing_compensation,
        "complete": missing_compensation == 0,
    }


def _style_report_sheet(ws, *, widths: list[float], freeze: str = "A4") -> None:
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    ws.auto_filter.ref = ws.dimensions
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.outlinePr.summaryBelow = True


def _daily_bridge_values(
    payload: dict[str, Any],
    products: list[dict[str, Any]],
) -> dict[str, float]:
    """Return one internally consistent basis for the daily profit bridge."""
    known_products = [
        item
        for item in products
        if item.get("cost_per_unit") is not None
        and float(item.get("known_cost_qty") or item.get("qty") or 0) > 0
    ]
    known_revenue = sum(
        float(
            item.get("known_revenue")
            if item.get("known_revenue") is not None
            else item.get("revenue") or 0
        )
        for item in known_products
    )
    known_payout = sum(
        float(
            item.get("known_payout")
            if item.get("known_payout") is not None
            else item.get("payout") or 0
        )
        for item in known_products
    )
    known_commission = sum(
        float(
            item.get("known_commission")
            if item.get("known_commission") is not None
            else item.get("commission") or 0
        )
        for item in known_products
    )
    known_logistics = sum(
        float(
            item.get("known_logistics")
            if item.get("known_logistics") is not None
            else item.get("logistics") or 0
        )
        for item in known_products
    )
    known_cost = sum(float(item.get("cost_total") or 0) for item in known_products)
    known_tax = sum(
        float(
            item.get("known_tax_expense")
            if item.get("known_tax_expense") is not None
            else item.get("tax_expense") or 0
        )
        for item in known_products
    )

    business = dict(payload.get("business") or {})
    revenue = float(business.get("calculation_revenue", known_revenue) or 0)
    payout = float(business.get("calculation_payout", known_payout) or 0)
    commission = float(
        business.get("calculation_commission", known_commission) or 0
    )
    logistics = float(business.get("calculation_logistics", known_logistics) or 0)
    residual = revenue - payout - commission - logistics
    other_payout = float(
        business.get("other_payout_deductions", max(0.0, residual)) or 0
    )
    payout_adjustment = float(
        business.get("payout_adjustment", max(0.0, -residual)) or 0
    )
    cost = float(business.get("cost_total", known_cost) or 0)
    tax = float(business.get("tax_expense", known_tax) or 0)
    deductions = max(0.0, float(payload.get("additional_expenses") or 0))
    refunds = max(0.0, float(payload.get("refunds") or 0))
    external = max(0.0, float(payload.get("external_expenses") or 0))
    profit_before_tax = payout - cost
    product_profit = profit_before_tax - tax
    final = product_profit - deductions + refunds - external
    return {
        "revenue": revenue,
        "commission": commission,
        "logistics": logistics,
        "other_payout": other_payout,
        "payout_adjustment": payout_adjustment,
        "payout": payout,
        "cost": cost,
        "profit_before_tax": profit_before_tax,
        "tax": tax,
        "product_profit": product_profit,
        "deductions": deductions,
        "refunds": refunds,
        "external": external,
        "final": final,
    }


def build_market_daily_workbook(
    payload: dict[str, Any],
    output: str | Path | None = None,
    *,
    lang: str = "uz",
) -> Path:
    path = _resolve_output(output, "sellerpro_daily_", ".xlsx")
    uz = str(lang).lower() == "uz"
    products = list(payload.get("products") or [])
    report_date = _date_label(payload.get("date"))
    shop_label = _clean(payload.get("shop_label") or payload.get("shop_id") or "—")

    wb = Workbook()
    ws = wb.active
    ws.title = "Kunlik hisobot" if uz else "Дневной отчёт"
    headers = (
        ["Tovar nomi", "Sxema", "Soni", "Daromad", "Chiqarishga", "Tannarx", "Soliq", "Sof foyda", "ROI %"]
        if uz
        else ["Товар", "Схема", "Кол-во", "Выручка", "К выплате", "Себестоимость", "Налог", "Чистая прибыль", "ROI %"]
    )

    ws.merge_cells("A1:I1")
    ws["A1"] = (f"Hisobot: {report_date}" if uz else f"Отчёт: {report_date}")
    ws["A1"].font = Font(name="Arial", size=16, bold=True, color=NAVY)
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 30
    ws.merge_cells("A2:I2")
    ws["A2"] = (f"Do‘kon: {shop_label}" if uz else f"Магазин: {shop_label}")
    ws["A2"].font = Font(name="Arial", size=10, color="566573")

    for column, header in enumerate(headers, start=1):
        cell = ws.cell(3, column, header)
        cell.font = Font(name="Arial", size=9, bold=True, color="333333")
        cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=GRID)
    ws.row_dimensions[3].height = 28

    first_data_row = 4
    for row_index, product in enumerate(products, start=first_data_row):
        has_cost = product.get("cost_per_unit") is not None and float(product.get("known_cost_qty") or product.get("qty") or 0) > 0
        sku_text = _identifier(product.get("sku"))
        values = [
            _clean(product.get("title") or sku_text or "—")
            + (f" - {sku_text}" if sku_text else ""),
            _clean(product.get("scheme") or "—"),
            _num(product.get("qty")) or 0,
            _num(product.get("revenue")) or 0,
            _num(product.get("payout")) or 0,
            _num(product.get("cost_total")) if has_cost else None,
            _num(product.get("tax_expense")) or 0,
            None,
            None,
        ]
        for column, value in enumerate(values, start=1):
            cell = ws.cell(row_index, column, value)
            cell.font = Font(name="Arial", size=8.5)
            cell.fill = PatternFill("solid", fgColor=WHITE if row_index % 2 == 0 else PALE_BLUE)
            cell.alignment = Alignment(
                horizontal="left" if column == 1 else "center",
                vertical="center",
                wrap_text=True,
            )
            cell.border = Border(bottom=GRID)
        if values[5] is not None:
            ws.cell(row_index, 8, f"=E{row_index}-F{row_index}-G{row_index}")
            ws.cell(row_index, 9, f'=IFERROR(H{row_index}/F{row_index},"")')
        ws.row_dimensions[row_index].height = 34

    last_data_row = max(first_data_row, first_data_row + len(products) - 1)
    total_row = last_data_row + 1
    ws.cell(total_row, 1, "Umumiy:" if uz else "Итого:")
    for column in range(3, 9):
        letter = get_column_letter(column)
        if products:
            ws.cell(total_row, column, f"=SUM({letter}{first_data_row}:{letter}{last_data_row})")
        else:
            ws.cell(total_row, column, 0)
    ws.cell(total_row, 9, f'=IFERROR(H{total_row}/F{total_row},"")')
    for cell in ws[total_row]:
        cell.font = Font(name="Arial", size=9, bold=True)
        cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=GRID, bottom=GRID)

    bridge = _daily_bridge_values(payload, products)
    row = total_row + 2
    summary_rows = [
        ("Hisobdagi tushum:" if uz else "Выручка в расчёте:", bridge["revenue"], NAVY, False),
        ("Uzum komissiyasi:" if uz else "Комиссия Uzum:", -bridge["commission"], RED, False),
        ("Logistika:" if uz else "Логистика:", -bridge["logistics"], RED, False),
        ("To‘lov ichidagi boshqa ushlanmalar:" if uz else "Другие удержания внутри выплаты:", -bridge["other_payout"], RED, False),
        ("To‘lov tuzatishi:" if uz else "Корректировка выплаты:", bridge["payout_adjustment"], GREEN, False),
        ("To‘lovga:" if uz else "К выплате:", None, NAVY, True),
        ("Tannarx:" if uz else "Себестоимость:", -bridge["cost"], RED, False),
        ("Soliqdan oldingi foyda:" if uz else "Прибыль до налога:", None, GREEN, True),
        ("Soliq:" if uz else "Налог:", -bridge["tax"], RED, False),
        ("Tovarlar foydasi:" if uz else "Прибыль товаров:", None, GREEN, True),
        ("Uzum qo‘shimcha xarajatlari:" if uz else "Доп. расходы Uzum:", -bridge["deductions"], RED, False),
        ("Uzum qaytargan mablag‘:" if uz else "Возвраты от Uzum:", bridge["refunds"], GREEN, False),
        ("Tashqi xarajatlar:" if uz else "Внешние расходы:", -bridge["external"], RED, False),
    ]
    summary_start = row
    for offset, (label, value, color, subtotal) in enumerate(summary_rows):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
        ws.cell(row, 1, label)
        if offset == 5:
            ws.cell(row, 8, f"=SUM(H{summary_start}:H{summary_start + 4})")
        elif offset == 7:
            ws.cell(row, 8, f"=H{summary_start + 5}+H{summary_start + 6}")
        elif offset == 9:
            ws.cell(row, 8, f"=H{summary_start + 7}+H{summary_start + 8}")
        else:
            ws.cell(row, 8, value)
        ws.cell(row, 8).font = Font(name="Arial", bold=True, color=color)
        ws.cell(row, 1).font = Font(name="Arial", bold=True, color=BLUE)
        if subtotal:
            for column in range(1, 9):
                ws.cell(row, column).border = Border(top=GRID)
        row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    complete = bool(payload.get("complete"))
    ws.cell(row, 1, "Yakuniy sof summa:" if uz and complete else "Ma’lum ma’lumotlar bo‘yicha natija:" if uz else "Итоговая чистая сумма:" if complete else "Результат по известным данным:")
    ws.cell(row, 8, f"=SUM(H{summary_start + 9}:H{summary_start + 12})")
    for cell in ws[row]:
        cell.font = Font(name="Arial", size=11, bold=True)
        cell.border = Border(top=Side(style="medium", color=BLUE))
    ws.cell(row, 8).alignment = Alignment(horizontal="right")
    row += 1
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=9)
    coverage = float(payload.get("cost_coverage") or 0) * 100
    note = (
        f"Tannarx qamrovi: {coverage:.1f}%. Manba: faqat Uzum purchasePrice."
        if uz
        else f"Покрытие себестоимостью: {coverage:.1f}%. Источник: только Uzum purchasePrice."
    )
    if not complete:
        note += (
            " Yetishmagan tannarxlar nol bilan almashtirilmagan."
            if uz
            else " Отсутствующая себестоимость не подставлена нулём."
        )
    ws.cell(row, 1, note)
    ws.cell(row, 1).font = Font(name="Arial", size=9, italic=True, color="6B7280")
    ws.cell(row, 1).alignment = Alignment(wrap_text=True)

    for column in ("D", "E", "F", "G", "H"):
        for cell in ws[column][first_data_row - 1 : row]:
            cell.number_format = '#,##0'
    for cell in ws["I"][first_data_row - 1 : total_row]:
        cell.number_format = "0.00%"
    _style_report_sheet(ws, widths=[74, 13, 9, 16, 16, 16, 14, 18, 12])
    ws.print_title_rows = "3:3"
    ws.auto_filter.ref = f"A3:I{last_data_row}" if products else "A3:I3"
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.save(path)
    wb.close()
    return path


def _pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    app_dir = Path(__file__).resolve().parent
    regular_candidates = [
        app_dir / "DejaVuSans.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
    ]
    bold_candidates = [
        app_dir / "DejaVuSans-Bold.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
    ]
    regular_path = next((item for item in regular_candidates if item.exists()), None)
    bold_path = next((item for item in bold_candidates if item.exists()), regular_path)
    if regular_path is None or bold_path is None:
        raise RuntimeError("Unicode PDF font was not found")
    if "SellerProDaily" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("SellerProDaily", str(regular_path)))
    if "SellerProDailyBold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("SellerProDailyBold", str(bold_path)))
    return "SellerProDaily", "SellerProDailyBold"


def build_market_daily_pdf(
    payload: dict[str, Any],
    output: str | Path | None = None,
    *,
    lang: str = "uz",
) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    path = _resolve_output(output, "sellerpro_daily_", ".pdf")
    regular, bold = _pdf_fonts()
    uz = str(lang).lower() == "uz"
    products = list(payload.get("products") or [])
    report_date = _date_label(payload.get("date"))
    shop_label = _clean(payload.get("shop_label") or payload.get("shop_id") or "—")
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title=f"SellerPro daily report {report_date}",
    )
    title_style = ParagraphStyle("DailyTitle", fontName=bold, fontSize=15, leading=18, textColor=colors.HexColor("#" + NAVY), alignment=TA_LEFT)
    meta_style = ParagraphStyle("DailyMeta", fontName=regular, fontSize=8, leading=10, textColor=colors.HexColor("#566573"))
    cell_style = ParagraphStyle("DailyCell", fontName=regular, fontSize=6.4, leading=8)
    cell_center = ParagraphStyle("DailyCellCenter", parent=cell_style, alignment=1)
    header_style = ParagraphStyle("DailyHeader", parent=cell_center, fontName=bold, fontSize=6.4, textColor=colors.HexColor("#333333"))
    story: list[Any] = [
        Paragraph((f"Hisobot: {report_date}" if uz else f"Отчёт: {report_date}"), title_style),
        Paragraph((f"Do‘kon: {xml_escape(shop_label)}" if uz else f"Магазин: {xml_escape(shop_label)}"), meta_style),
        Spacer(1, 4 * mm),
    ]
    headers = (
        ["Tovar nomi", "Sxema", "Soni", "Daromad", "Chiqarishga", "Tannarx", "Soliq", "Sof foyda", "ROI %"]
        if uz
        else ["Товар", "Схема", "Кол-во", "Выручка", "К выплате", "Себестоимость", "Налог", "Чистая прибыль", "ROI %"]
    )
    header_row: list[Any] = [Paragraph(value, header_style) for value in headers]
    product_rows: list[list[Any]] = []
    for product in products:
        title = _clean(product.get("title") or "—")
        sku = _identifier(product.get("sku"))
        if sku:
            title += f" - {sku}"
        has_cost = product.get("cost_per_unit") is not None and float(product.get("known_cost_qty") or product.get("qty") or 0) > 0
        cost = _num(product.get("cost_total")) if has_cost else None
        tax = _num(product.get("tax_expense")) or 0.0
        payout = _num(product.get("payout")) or 0.0
        net = payout - cost - tax if cost is not None else None
        roi = net / cost * 100.0 if net is not None and cost and cost > 0 else None
        product_rows.append(
            [
                Paragraph(xml_escape(title), cell_style),
                Paragraph(xml_escape(_clean(product.get("scheme") or "—")), cell_center),
                Paragraph(f"{float(product.get('qty') or 0):g}", cell_center),
                Paragraph(_money(product.get("revenue")), cell_center),
                Paragraph(_money(payout), cell_center),
                Paragraph(_money(cost), cell_center),
                Paragraph(_money(tax), cell_center),
                Paragraph(_money(net), cell_center),
                Paragraph("—" if roi is None else f"{roi:.2f}%", cell_center),
            ]
        )
    totals = {
        key: sum(float(item.get(key) or 0) for item in products)
        for key in ("qty", "revenue", "payout", "cost_total", "tax_expense")
    }
    known_products = [
        item
        for item in products
        if item.get("cost_per_unit") is not None
        and float(item.get("known_cost_qty") or item.get("qty") or 0) > 0
    ]
    known_revenue = sum(
        float(item.get("known_revenue") if item.get("known_revenue") is not None else item.get("revenue") or 0)
        for item in known_products
    )
    known_payout = sum(
        float(item.get("known_payout") if item.get("known_payout") is not None else item.get("payout") or 0)
        for item in known_products
    )
    known_tax = sum(
        float(item.get("known_tax_expense") if item.get("known_tax_expense") is not None else item.get("tax_expense") or 0)
        for item in known_products
    )
    known_net = sum(
        float(item.get("known_payout") if item.get("known_payout") is not None else item.get("payout") or 0)
        - float(item.get("cost_total") or 0)
        - float(item.get("known_tax_expense") if item.get("known_tax_expense") is not None else item.get("tax_expense") or 0)
        for item in known_products
    )
    roi_total = known_net / totals["cost_total"] * 100.0 if totals["cost_total"] > 0 else None
    total_pdf_row = [
        Paragraph("Umumiy:" if uz else "Итого:", header_style),
        "",
        Paragraph(f"{totals['qty']:g}", header_style),
        Paragraph(_money(totals["revenue"]), header_style),
        Paragraph(_money(totals["payout"]), header_style),
        Paragraph(_money(totals["cost_total"]), header_style),
        Paragraph(_money(totals["tax_expense"]), header_style),
        Paragraph(_money(known_net), header_style),
        Paragraph("—" if roi_total is None else f"{roi_total:.2f}%", header_style),
    ]
    # Fixed-size chunks avoid leaving one orphan product row next to the final
    # summary.  Each chunk is a separate table, so its header is guaranteed.
    page_size = 21
    chunks = [product_rows[index : index + page_size] for index in range(0, len(product_rows), page_size)] or [[]]
    widths = [103 * mm, 15 * mm, 12 * mm, 25 * mm, 25 * mm, 25 * mm, 21 * mm, 26 * mm, 18 * mm]
    for chunk_index, chunk in enumerate(chunks):
        last_chunk = chunk_index == len(chunks) - 1
        data = [header_row, *chunk]
        if last_chunk:
            data.append(total_pdf_row)
        table = Table(data, colWidths=widths, repeatRows=1)
        body_end = -2 if last_chunk else -1
        styles = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + PALE_GRAY)),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EA")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        if chunk:
            styles.append(("ROWBACKGROUNDS", (0, 1), (-1, body_end), [colors.white, colors.HexColor("#" + PALE_BLUE)]))
        if last_chunk:
            styles.append(("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#" + PALE_GRAY)))
        table.setStyle(TableStyle(styles))
        story.append(table)
        if last_chunk:
            story.append(Spacer(1, 4 * mm))
        else:
            story.append(PageBreak())
    bridge = _daily_bridge_values(payload, products)
    final = bridge["final"]
    complete = bool(payload.get("complete"))
    coverage = float(payload.get("cost_coverage") or 0) * 100
    summary_data = [
        ["Hisobdagi tushum" if uz else "Выручка в расчёте", _signed_money(bridge["revenue"], "+")],
        ["Uzum komissiyasi" if uz else "Комиссия Uzum", _signed_money(bridge["commission"], "-")],
        ["Logistika", _signed_money(bridge["logistics"], "-")],
        ["To‘lov ichidagi boshqa ushlanmalar" if uz else "Другие удержания внутри выплаты", _signed_money(bridge["other_payout"], "-")],
        ["To‘lov tuzatishi" if uz else "Корректировка выплаты", _signed_money(bridge["payout_adjustment"], "+")],
        ["To‘lovga" if uz else "К выплате", _money(bridge["payout"])],
        ["Tannarx" if uz else "Себестоимость", _signed_money(bridge["cost"], "-")],
        ["Soliqdan oldingi foyda" if uz else "Прибыль до налога", _money(bridge["profit_before_tax"])],
        ["Soliq" if uz else "Налог", _signed_money(bridge["tax"], "-")],
        ["Tovarlar foydasi" if uz else "Прибыль товаров", _money(bridge["product_profit"])],
        ["Uzum qo‘shimcha xarajatlari" if uz else "Доп. расходы Uzum", _signed_money(bridge["deductions"], "-")],
        ["Uzum qaytargan mablag‘" if uz else "Возвраты от Uzum", _signed_money(bridge["refunds"], "+")],
        ["Tashqi xarajatlar" if uz else "Внешние расходы", _signed_money(bridge["external"], "-")],
        [
            "Yakuniy sof summa" if uz and complete else "Ma’lum ma’lumotlar bo‘yicha natija" if uz else "Итоговая чистая сумма" if complete else "Результат по известным данным",
            _money(final),
        ],
        ["ROI %", "—" if bridge["cost"] <= 0 else f"{final / bridge['cost'] * 100:.2f}%"],
    ]
    summary_table = Table(summary_data, colWidths=[80 * mm, 40 * mm], hAlign="RIGHT")
    summary_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), regular),
                ("FONTNAME", (0, 5), (-1, 5), bold),
                ("FONTNAME", (0, 7), (-1, 7), bold),
                ("FONTNAME", (0, 9), (-1, 9), bold),
                ("FONTNAME", (0, -2), (-1, -1), bold),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("LINEABOVE", (0, 5), (-1, 5), 0.5, colors.HexColor("#" + BLUE)),
                ("LINEABOVE", (0, 7), (-1, 7), 0.5, colors.HexColor("#" + BLUE)),
                ("LINEABOVE", (0, 9), (-1, 9), 0.5, colors.HexColor("#" + BLUE)),
                ("LINEABOVE", (0, -2), (-1, -2), 1, colors.HexColor("#" + BLUE)),
                ("TEXTCOLOR", (1, 1), (1, 3), colors.HexColor("#" + RED)),
                ("TEXTCOLOR", (1, 4), (1, 4), colors.HexColor("#" + GREEN)),
                ("TEXTCOLOR", (1, 6), (1, 6), colors.HexColor("#" + RED)),
                ("TEXTCOLOR", (1, 8), (1, 8), colors.HexColor("#" + RED)),
                ("TEXTCOLOR", (1, 10), (1, 10), colors.HexColor("#" + RED)),
                ("TEXTCOLOR", (1, 11), (1, 11), colors.HexColor("#" + GREEN)),
                ("TEXTCOLOR", (1, 12), (1, 12), colors.HexColor("#" + RED)),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    note = (
        f"Tannarx qamrovi: {coverage:.1f}%. Manba: faqat Uzum purchasePrice."
        if uz
        else f"Покрытие себестоимостью: {coverage:.1f}%. Источник: только Uzum purchasePrice."
    )
    story.extend([KeepTogether([summary_table, Spacer(1, 2 * mm), Paragraph(note, meta_style)])])
    doc.build(story)
    return path


def build_compensation_workbook(
    rows: list[dict[str, Any]],
    output: str | Path | None = None,
    *,
    kind: str,
    lang: str = "ru",
) -> Path:
    normalized_kind = str(kind).lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    path = _resolve_output(output, f"sellerpro_{normalized_kind}_", ".xlsx")
    uz = str(lang).lower() == "uz"
    wb = Workbook()
    ws = wb.active
    ws.title = (
        "Yo‘qolgan tovarlar" if uz and normalized_kind == "loss" else
        "Brak hisoboti" if uz else
        "Потерянные товары" if normalized_kind == "loss" else
        "Отчет браков"
    )
    if normalized_kind == "loss":
        headers = (
            ["Tovar nomi", "Shtrix-kod", "Qoplama miqdori", "Soni", "Jami"]
            if uz
            else ["Название товара", "Штрих-код", "Размер возмещения", "Количество", "Итого"]
        )
        widths = [82, 22, 22, 14, 22]
    else:
        headers = (
            ["Tovar ID", "Tovar nomi", "SKU", "Shtrix-kod", "Qoplama miqdori", "Soni", "Jami"]
            if uz
            else ["ИД товара", "Название товара", "SKU", "ШК", "Размер возмещения", "Кол-во", "Итого"]
        )
        widths = [16, 65, 24, 22, 22, 12, 22]
    total_col = len(headers)
    total_qty = sum(_qty(item.get("quantity")) for item in rows)
    known_total = sum(float(_num(item.get("total")) or 0) for item in rows)
    incomplete = sum(
        1
        for item in rows
        if _num(item.get("compensation")) is None or _num(item.get("total")) is None
    )
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_col)
    ws.cell(1, 1, "Yo‘qolgan tovarlar reestri" if uz and normalized_kind == "loss" else "Brak tovarlar reestri" if uz else "Реестр потерянных товаров" if normalized_kind == "loss" else "Реестр повреждённых товаров")
    ws.cell(1, 1).font = Font(name="Arial", size=16, bold=True, color=NAVY)
    ws.cell(1, 1).alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 30
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_col)
    ws.cell(2, 1, (
        f"Pozitsiyalar: {len(rows)} · Jami birlik: {total_qty} · Hisoblangan summa: {_money(known_total)}"
        if uz
        else f"Позиций: {len(rows)} · Единиц: {total_qty} · Рассчитанная сумма: {_money(known_total)}"
    ))
    ws.cell(2, 1).font = Font(name="Arial", size=10, bold=True, color="34495E")
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=total_col)
    ws.cell(3, 1, (
        f"Diqqat: {incomplete} qator bo‘yicha aniq narx yoki komissiya yo‘q; summa bo‘sh qoldirildi."
        if uz and incomplete
        else "Ma’lumotlar Uzum API asosida. Summalar taxmin qilinmaydi."
        if uz
        else f"Внимание: по {incomplete} строкам нет точной цены или комиссии; сумма оставлена пустой."
        if incomplete
        else "Данные получены из Uzum API. Отсутствующие суммы не подставляются приблизительно."
    ))
    ws.cell(3, 1).font = Font(name="Arial", size=9, italic=True, color=RED if incomplete else "6B7280")
    header_row = 4
    for column, header in enumerate(headers, start=1):
        ws.cell(header_row, column, header)
    for cell in ws[header_row]:
        cell.font = Font(name="Arial", bold=True, color="333333")
        cell.fill = PatternFill("solid", fgColor="E7E7E7")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=GRID)
    ws.row_dimensions[header_row].height = 28

    for index, item in enumerate(rows, start=header_row + 1):
        compensation = _num(item.get("compensation"))
        if normalized_kind == "loss":
            values = [item.get("title"), item.get("barcode"), compensation, item.get("quantity"), item.get("total")]
        else:
            values = [item.get("product_id"), item.get("title"), item.get("sku_id") or item.get("sku"), item.get("barcode"), compensation, item.get("quantity"), item.get("total")]
        ws.append(values)
        for cell in ws[index]:
            cell.fill = PatternFill("solid", fgColor=WHITE if index % 2 == 0 else PALE_BLUE)
            cell.font = Font(name="Arial", size=9)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(bottom=GRID)
        ws.row_dimensions[index].height = 34

    total_row = header_row + len(rows) + 1
    qty_col = total_col - 1
    ws.cell(total_row, 1, "Umumiy:" if uz else "Итого:")
    ws.cell(total_row, qty_col, f"=SUM({get_column_letter(qty_col)}{header_row + 1}:{get_column_letter(qty_col)}{total_row - 1})")
    ws.cell(total_row, total_col, f"=SUM({get_column_letter(total_col)}{header_row + 1}:{get_column_letter(total_col)}{total_row - 1})")
    for cell in ws[total_row]:
        cell.font = Font(name="Arial", bold=True)
        cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
        cell.border = Border(top=GRID, bottom=GRID)
    for column in (total_col - 2, total_col):
        for cell in ws[get_column_letter(column)][header_row:]:
            cell.number_format = '#,##0'
    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(total_col)}{max(header_row, total_row - 1)}"
    ws.sheet_view.showGridLines = False
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.page_setup.orientation = "landscape"
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.print_title_rows = f"{header_row}:{header_row}"
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.save(path)
    wb.close()
    return path


def _set_docx_run_font(run: Any, *, bold: bool | None = None, size: float | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def build_claim_docx(
    rows: list[dict[str, Any]],
    output: str | Path | None = None,
    *,
    kind: str,
    shop_id: int | str | None = None,
    seller_details: dict[str, Any] | None = None,
    lang: str = "ru",
) -> Path:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    normalized_kind = str(kind).lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    path = _resolve_output(output, f"sellerpro_claim_{normalized_kind}_", ".docx")
    uz = str(lang).lower() == "uz"
    details = dict(seller_details or {})
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark = header.add_run("⬡ ")
    _set_docx_run_font(mark, bold=True, size=15)
    mark.font.color.rgb = RGBColor(108, 49, 217)
    brand = header.add_run("UZUM MARKET")
    _set_docx_run_font(brand, bold=True, size=13)
    brand.font.color.rgb = RGBColor(108, 49, 217)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("PRETENZIYA" if uz else "ПРЕТЕНЗИЯ")
    _set_docx_run_font(run, bold=True, size=14)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(5)
    intro_text = (
        "Quyidagi tovarlar bo‘yicha kompensatsiya to‘lashni so‘rayman. "
        "Maydonlarni yuborishdan oldin sotuvchi ma’lumotlari bilan to‘ldiring."
        if uz
        else "Прошу выплатить компенсацию по перечисленным ниже товарам. "
        "Перед отправкой заполните реквизиты продавца."
    )
    run = intro.add_run(intro_text)
    _set_docx_run_font(run, size=8.5)

    labels = (
        [
            ("Sotuvchi F.I.Sh.", "full_name"),
            ("YaTT / tashkilot", "company"),
            ("JShShIR / STIR", "tax_id"),
            ("Hisob raqami", "bank_account"),
            ("MFO", "mfo"),
            ("Do‘kon ID", "shop_id"),
        ]
        if uz
        else [
            ("Ф.И.О. продавца", "full_name"),
            ("ИП / организация", "company"),
            ("ПИНФЛ / ИНН", "tax_id"),
            ("Расчётный счёт", "bank_account"),
            ("МФО", "mfo"),
            ("ID магазина", "shop_id"),
        ]
    )
    details["shop_id"] = details.get("shop_id") or shop_id or ""
    details_table = document.add_table(rows=3, cols=4)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False
    widths = [Cm(3.2), Cm(6.2), Cm(3.0), Cm(6.0)]
    for index, (label, key) in enumerate(labels):
        table_row = index // 2
        pair = (index % 2) * 2
        details_table.cell(table_row, pair).width = widths[pair]
        details_table.cell(table_row, pair + 1).width = widths[pair + 1]
        details_table.cell(table_row, pair).text = label
        details_table.cell(table_row, pair + 1).text = _clean(details.get(key)) or "____________________________"
    for row in details_table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for item in paragraph.runs:
                    _set_docx_run_font(item, bold=index % 2 == 0, size=7.5)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = "w:" + edge
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "4")
                element.set(qn("w:color"), "D9E2EA")

    preview_qty = sum(_qty(item.get("quantity")) for item in rows)
    preview_total = sum(float(_num(item.get("total")) or 0) for item in rows)
    preview_incomplete = sum(
        1
        for item in rows
        if _num(item.get("compensation")) is None or _num(item.get("total")) is None
    )
    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(5)
    summary.paragraph_format.space_after = Pt(5)
    summary_text = (
        f"Hujjat tarkibi: {len(rows)} pozitsiya · {preview_qty} dona · hisoblangan summa {_money(preview_total)} so‘m."
        if uz
        else f"Состав претензии: {len(rows)} позиций · {preview_qty} единиц · рассчитанная сумма {_money(preview_total)} сум."
    )
    if preview_incomplete:
        summary_text += (
            f" {preview_incomplete} qator bo‘yicha summa ma’lumot yetishmagani sababli bo‘sh qoldirilgan."
            if uz
            else f" По {preview_incomplete} строкам сумма оставлена пустой из-за отсутствующих данных."
        )
    summary_run = summary.add_run(summary_text)
    _set_docx_run_font(summary_run, bold=True, size=8.2)
    summary_run.font.color.rgb = RGBColor(23, 54, 93)

    reason = "Yo‘qotish" if uz and normalized_kind == "loss" else "Shikastlanish" if uz else "Утеря" if normalized_kind == "loss" else "Повреждение"
    headers = (
        ["Tovar nomi", "Shtrix-kod", "Sabab", "Qoplama*", "Soni", "Jami"]
        if uz
        else ["Название товара", "Штрих-код", "Причина", "Размер возмещения*", "Кол-во", "Итого"]
    )
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_widths = [Cm(8.0), Cm(2.7), Cm(2.2), Cm(2.6), Cm(1.2), Cm(2.5)]
    for index, header_text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = column_widths[index]
        cell.text = header_text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E7E7E7")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item in paragraph.runs:
                _set_docx_run_font(item, bold=True, size=6.6)

    known_total = 0.0
    total_qty = 0
    incomplete = 0
    for index, item in enumerate(rows, start=1):
        row = table.add_row()
        compensation = _num(item.get("compensation"))
        total = _num(item.get("total"))
        if compensation is None or total is None:
            incomplete += 1
        else:
            known_total += total
        quantity = _qty(item.get("quantity"))
        total_qty += quantity
        title_text = _clean(item.get("title"))
        sku_text = _identifier(item.get("sku"))
        if sku_text:
            title_text += f" - {sku_text}"
        values = [
            title_text,
            _identifier(item.get("barcode")) or "—",
            reason,
            _money(compensation),
            str(quantity),
            _money(total),
        ]
        for column, value in enumerate(values):
            cell = row.cells[column]
            cell.width = column_widths[column]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if index % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), PALE_BLUE)
                cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for item_run in paragraph.runs:
                    _set_docx_run_font(item_run, size=6.2)
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)

    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = "Umumiy:" if uz else "Итого:"
    total_row.cells[4].text = str(total_qty)
    total_row.cells[5].text = _money(known_total)
    for cell in total_row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E7E7E7")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item in paragraph.runs:
                _set_docx_run_font(item, bold=True, size=7)

    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().insert(0, table_header)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(2)
    note_text = (
        "* Qoplama miqdori Uzumdagi joriy sotuv narxidan SKU komissiyasi ayrilgan holda hisoblandi."
        if uz
        else "* Размер возмещения рассчитан как текущая стоимость товара в Uzum за вычетом комиссии маркетплейса по SKU."
    )
    if incomplete:
        note_text += (
            f" {incomplete} ta SKU bo‘yicha narx yoki aniq komissiya yo‘q; qiymat taxmin qilinmadi."
            if uz
            else f" По {incomplete} SKU нет цены или точной комиссии; сумма не рассчитана и оставлена пустой."
        )
    run = note.add_run(note_text)
    _set_docx_run_font(run, size=7.3)
    run.italic = True

    signature = document.add_paragraph()
    signature.paragraph_format.space_before = Pt(8)
    signature.add_run(("Sana: ______________    Imzo: ____________________" if uz else "Дата: ______________    Подпись: ____________________"))
    for item in signature.runs:
        _set_docx_run_font(item, size=8.5)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run("Seller.pro.uz · Uzum API ma’lumotlari" if uz else "Seller.pro.uz · данные Uzum API")
    _set_docx_run_font(footer_run, size=7)
    footer_run.font.color.rgb = RGBColor(128, 139, 150)

    document.save(path)
    return path
