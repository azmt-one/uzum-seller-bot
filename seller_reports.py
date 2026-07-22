from __future__ import annotations

import math
import tempfile
from datetime import date, datetime
from html import escape as xml_escape
from pathlib import Path
from typing import Any, Iterable

from openpyxl import Workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


NAVY = "17365D"
PURPLE = "6C31D9"
BLUE = "1F9AD6"
PALE_BLUE = "EAF4FA"
PALE_GRAY = "F5F7FA"
GREEN = "2E9B4F"
RED = "C0392B"
WHITE = "FFFFFF"
GRID = Side(style="thin", color="D9E2EA")


def _num(value: Any) -> float | None:
    if value is None or isinstance(value, bool):
        return None
    try:
        result = float(value)
    except (TypeError, ValueError):
        return None
    return result if math.isfinite(result) else None


def _money(value: Any) -> str:
    number = _num(value)
    return "—" if number is None else f"{number:,.0f}".replace(",", " ")


def _signed_money(value: Any, sign: str) -> str:
    number = max(0.0, float(_num(value) or 0))
    return _money(0) if number <= 0 else f"{sign} {_money(number)}"


def _qty(value: Any) -> int:
    number = _num(value)
    return max(0, int(number or 0))


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _identifier(value: Any) -> str:
    if value is None or isinstance(value, bool):
        return ""
    if isinstance(value, (int, float)):
        number = float(value)
        if math.isfinite(number) and number.is_integer():
            return str(int(number))
    text = _clean(value)
    if text.endswith(".0"):
        try:
            number = float(text)
            if number.is_integer():
                return str(int(number))
        except ValueError:
            pass
    return text


def _date_label(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return _clean(value) or "—"


def _resolve_output(output: str | Path | None, prefix: str, suffix: str) -> Path:
    if output is None:
        handle = tempfile.NamedTemporaryFile(prefix=prefix, suffix=suffix, delete=False)
        handle.close()
        return Path(handle.name)
    path = Path(output)
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def prepare_compensation_rows(
    stock_rows: Iterable[dict[str, Any]],
    *,
    kind: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Build claim rows only from documented Uzum catalog values.

    Compensation is calculated as the current product value less the exact
    marketplace commission.  We calculate it only when both the sale price and
    exact SKU commission are present.  A commission range is never guessed.
    """

    normalized_kind = str(kind).strip().lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    quantity_field = "missing" if normalized_kind == "loss" else "defected"
    rows: list[dict[str, Any]] = []
    missing_compensation = 0
    total_qty = 0
    total_value = 0.0

    for source in stock_rows:
        qty = _qty(source.get(quantity_field))
        if qty <= 0:
            continue
        price = _num(source.get("price"))
        commission = _num(source.get("commission"))
        compensation: float | None = None
        if (
            price is not None
            and price > 0
            and commission is not None
            and 0 <= commission <= 100
        ):
            compensation = max(0.0, price * (1.0 - commission / 100.0))
        else:
            missing_compensation += 1

        total = compensation * qty if compensation is not None else None
        total_qty += qty
        if total is not None:
            total_value += total
        rows.append(
            {
                "product_id": _identifier(source.get("product_id")),
                "title": _clean(
                    source.get("product_title")
                    or source.get("sku_full_title")
                    or source.get("sku_title")
                    or "Без названия"
                ),
                "sku": _clean(
                    source.get("seller_item_code")
                    or source.get("sku_full_title")
                    or source.get("sku_title")
                    or source.get("sku_id")
                ),
                "sku_id": _identifier(source.get("sku_id")),
                "barcode": _identifier(source.get("barcode")),
                "price": price,
                "commission_percent": commission,
                "compensation": compensation,
                "quantity": qty,
                "total": total,
            }
        )

    rows.sort(key=lambda item: (-int(item["quantity"]), str(item["title"]).lower()))
    return rows, {
        "kind": normalized_kind,
        "rows": len(rows),
        "quantity": total_qty,
        "known_total": total_value,
        "missing_compensation": missing_compensation,
        "complete": missing_compensation == 0,
    }


def _style_report_sheet(ws, *, widths: list[float], freeze: str = "A4") -> None:
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.outlinePr.summaryBelow = True


def _daily_bridge_values(
    payload: dict[str, Any],
    products: list[dict[str, Any]],
) -> dict[str, float]:
    """Return one internally consistent basis for the daily profit bridge."""
    known_products = [
        item
        for item in products
        if item.get("cost_per_unit") is not None
        and float(item.get("known_cost_qty") or item.get("qty") or 0) > 0
    ]
    known_revenue = sum(
        float(
            item.get("known_revenue")
            if item.get("known_revenue") is not None
            else item.get("revenue") or 0
        )
        for item in known_products
    )
    known_payout = sum(
        float(
            item.get("known_payout")
            if item.get("known_payout") is not None
            else item.get("payout") or 0
        )
        for item in known_products
    )
    known_commission = sum(
        float(
            item.get("known_commission")
            if item.get("known_commission") is not None
            else item.get("commission") or 0
        )
        for item in known_products
    )
    known_logistics = sum(
        float(
            item.get("known_logistics")
            if item.get("known_logistics") is not None
            else item.get("logistics") or 0
        )
        for item in known_products
    )
    known_cost = sum(float(item.get("cost_total") or 0) for item in known_products)
    known_tax = sum(
        float(
            item.get("known_tax_expense")
            if item.get("known_tax_expense") is not None
            else item.get("tax_expense") or 0
        )
        for item in known_products
    )

    business = dict(payload.get("business") or {})
    revenue = float(business.get("calculation_revenue", known_revenue) or 0)
    payout = float(business.get("calculation_payout", known_payout) or 0)
    commission = float(
        business.get("calculation_commission", known_commission) or 0
    )
    logistics = float(business.get("calculation_logistics", known_logistics) or 0)
    residual = revenue - payout - commission - logistics
    other_payout = float(
        business.get("other_payout_deductions", max(0.0, residual)) or 0
    )
    payout_adjustment = float(
        business.get("payout_adjustment", max(0.0, -residual)) or 0
    )
    cost = float(business.get("cost_total", known_cost) or 0)
    tax = float(business.get("tax_expense", known_tax) or 0)
    deductions = max(0.0, float(payload.get("additional_expenses") or 0))
    refunds = max(0.0, float(payload.get("refunds") or 0))
    external = max(0.0, float(payload.get("external_expenses") or 0))
    profit_before_tax = payout - cost
    product_profit = profit_before_tax - tax
    final = product_profit - deductions + refunds - external
    return {
        "revenue": revenue,
        "commission": commission,
        "logistics": logistics,
        "other_payout": other_payout,
        "payout_adjustment": payout_adjustment,
        "payout": payout,
        "cost": cost,
        "profit_before_tax": profit_before_tax,
        "tax": tax,
        "product_profit": product_profit,
        "deductions": deductions,
        "refunds": refunds,
        "external": external,
        "final": final,
    }


def build_daily_finance_workbook(
    payload: dict[str, Any],
    output: str | Path | None = None,
    *,
    lang: str = "uz",
) -> Path:
    path = _resolve_output(output, "sellerpro_daily_", ".xlsx")
    uz = str(lang).lower() == "uz"
    products = list(payload.get("products") or [])
    report_date = _date_label(payload.get("date"))
    shop_label = _clean(payload.get("shop_label") or payload.get("shop_id") or "—")
    bridge = _daily_bridge_values(payload, products)
    complete = bool(payload.get("complete"))
    coverage = float(payload.get("cost_coverage") or 0) * 100
    accepted = dict(payload.get("accepted") or {})
    issued = dict(payload.get("issued") or {})
    schemes = dict(payload.get("schemes") or {})
    expenses = sorted(
        (payload.get("expense_names") or {}).items(),
        key=lambda pair: abs(float(pair[1] or 0)),
        reverse=True,
    )
    result_label = (
        "Yakuniy sof foyda"
        if uz and complete
        else "Ma’lum ma’lumotlar bo‘yicha natija"
        if uz
        else "Итоговая чистая прибыль"
        if complete
        else "Результат по известным данным"
    )
    money_format = '#,##0;[Red]-#,##0;–'
    percent_format = '0.0%;[Red]-0.0%;–'

    wb = Workbook()
    ws = wb.active
    ws.title = "Xulosa" if uz else "Сводка"
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:D1")
    ws["A1"] = "Seller.pro.uz · Kunlik hisobot" if uz else "Seller.pro.uz · Дневной отчёт"
    ws["A1"].font = Font(name="Arial", size=17, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor=NAVY)
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 34
    ws.merge_cells("A2:D2")
    ws["A2"] = (
        f"Sana: {report_date} · Do‘kon: {shop_label} · Valyuta: so‘m"
        if uz
        else f"Дата: {report_date} · Магазин: {shop_label} · Валюта: сум"
    )
    ws["A2"].font = Font(name="Arial", size=10, color="566573")
    ws["A2"].alignment = Alignment(vertical="center")

    def section_title(row: int, text: str) -> None:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        cell = ws.cell(row, 1, text)
        cell.font = Font(name="Arial", size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(vertical="center")
        ws.row_dimensions[row].height = 23

    section_title(4, "Kun yakuni" if uz else "Главное за день")
    kpis = [
        ("Qabul qilingan buyurtmalar" if uz else "Принято заказов", int(accepted.get("orders") or 0)),
        ("Sotilgan tovarlar" if uz else "Продано товаров", float(accepted.get("units") or 0)),
        ("Tushum" if uz else "Выручка", bridge["revenue"]),
        ("To‘lovga" if uz else "К выплате", bridge["payout"]),
        (result_label, "=B24"),
        ("ROI", '=IFERROR(B24/-B17,"")'),
    ]
    for index, (label, value) in enumerate(kpis):
        row = 5 + index // 2
        label_col = 1 if index % 2 == 0 else 3
        value_col = label_col + 1
        ws.cell(row, label_col, label)
        ws.cell(row, value_col, value)
        ws.cell(row, label_col).font = Font(name="Arial", size=9, bold=True, color="566573")
        ws.cell(row, value_col).font = Font(name="Arial", size=11, bold=True, color=NAVY)
        ws.cell(row, label_col).fill = PatternFill("solid", fgColor=PALE_GRAY)
        ws.cell(row, value_col).fill = PatternFill("solid", fgColor=PALE_BLUE)
        for column in (label_col, value_col):
            ws.cell(row, column).border = Border(bottom=GRID)
            ws.cell(row, column).alignment = Alignment(vertical="center", wrap_text=True)
        ws.cell(row, value_col).alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 29
    for cell in (ws["B6"], ws["D6"], ws["B7"]):
        cell.number_format = money_format
    ws["D7"].number_format = percent_format
    final_fill = "E2F0D9" if bridge["final"] >= 0 else "FCE4D6"
    final_color = GREEN if bridge["final"] >= 0 else RED
    ws["B7"].fill = PatternFill("solid", fgColor=final_fill)
    ws["B7"].font = Font(name="Arial", size=12, bold=True, color=final_color)

    section_title(9, "Natija qanday hisoblandi" if uz else "Как получился итог")
    bridge_headers = (
        ["Qadam", "Summa", "Izoh", ""]
        if uz
        else ["Шаг", "Сумма", "Что означает", ""]
    )
    for column, value in enumerate(bridge_headers, start=1):
        ws.cell(10, column, value)
        ws.cell(10, column).font = Font(name="Arial", size=9, bold=True, color="333333")
        ws.cell(10, column).fill = PatternFill("solid", fgColor=PALE_GRAY)
        ws.cell(10, column).border = Border(bottom=GRID)
        ws.cell(10, column).alignment = Alignment(vertical="center")
    ws.merge_cells("C10:D10")

    bridge_rows = [
        ("1. Tushum" if uz else "1. Выручка", bridge["revenue"], "Sotilgan tovarlar narxi" if uz else "Цена проданных товаров", False),
        ("2. − Uzum komissiyasi" if uz else "2. − Комиссия Uzum", -bridge["commission"], "To‘lov summasida allaqachon hisobga olingan" if uz else "Уже учтена в сумме «К выплате»", False),
        ("3. − Logistika" if uz else "3. − Логистика", -bridge["logistics"], "To‘lov summasida allaqachon hisobga olingan" if uz else "Уже учтена в сумме «К выплате»", False),
        ("4. − Boshqa ushlanmalar" if uz else "4. − Другие удержания", -bridge["other_payout"], "To‘lov ichidagi boshqa ushlanmalar" if uz else "Другие удержания внутри выплаты", False),
        ("5. + To‘lov tuzatishi" if uz else "5. + Корректировка выплаты", bridge["payout_adjustment"], "Uzum bergan to‘lov bilan farq" if uz else "Разница с фактической выплатой Uzum", False),
        ("To‘lovga" if uz else "К выплате", "=SUM(B11:B15)", "Uzum ushlanmalaridan keyin" if uz else "После удержаний Uzum", True),
        ("6. − Tannarx" if uz else "6. − Себестоимость", -bridge["cost"], "Uzum purchasePrice qiymati" if uz else "Закупочная стоимость из Uzum purchasePrice", False),
        ("Soliqdan oldingi foyda" if uz else "Прибыль до налога", "=SUM(B16:B17)", "To‘lov minus tannarx" if uz else "К выплате минус себестоимость", True),
        ("7. − Soliq" if uz else "7. − Налог", -bridge["tax"], "Bot sozlamalaridagi stavka" if uz else "По ставке из настроек бота", False),
        ("Tovarlar foydasi" if uz else "Прибыль товаров", "=SUM(B18:B19)", "Qo‘shimcha xarajatlardan oldin" if uz else "До дополнительных расходов", True),
        ("8. − Uzum qo‘shimcha xarajatlari" if uz else "8. − Доп. расходы Uzum", -bridge["deductions"], "Saqlash, reklama va boshqa xizmatlar" if uz else "Хранение, продвижение и другие услуги", False),
        ("9. + Uzum qaytargan mablag‘" if uz else "9. + Возвраты Uzum", bridge["refunds"], "Uzum qaytargan summa" if uz else "Средства, возвращённые Uzum", False),
        ("10. − Tashqi xarajatlar" if uz else "10. − Внешние расходы", -bridge["external"], "Uzumdan tashqaridagi qo‘lda kiritilgan xarajatlar" if uz else "Расходы вне Uzum, добавленные вручную", False),
        (result_label, "=SUM(B20:B23)", "Barcha ma’lum xarajatlardan keyin" if uz else "После всех известных расходов", True),
    ]
    for row, (label, value, explanation, subtotal) in enumerate(bridge_rows, start=11):
        ws.cell(row, 1, label)
        ws.cell(row, 2, value)
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        ws.cell(row, 3, explanation)
        ws.cell(row, 1).font = Font(name="Arial", size=9, bold=subtotal)
        ws.cell(row, 2).font = Font(name="Arial", size=9, bold=subtotal)
        ws.cell(row, 3).font = Font(name="Arial", size=8.5, color="566573")
        ws.cell(row, 2).number_format = money_format
        ws.cell(row, 2).alignment = Alignment(horizontal="right", vertical="center")
        ws.cell(row, 3).alignment = Alignment(vertical="center", wrap_text=True)
        if subtotal:
            for column in range(1, 5):
                ws.cell(row, column).fill = PatternFill("solid", fgColor=PALE_BLUE)
                ws.cell(row, column).border = Border(top=GRID, bottom=GRID)
        else:
            for column in range(1, 5):
                ws.cell(row, column).border = Border(bottom=GRID)
        ws.row_dimensions[row].height = 25
    ws["B24"].fill = PatternFill("solid", fgColor=final_fill)
    ws["B24"].font = Font(name="Arial", size=11, bold=True, color=final_color)

    ws.merge_cells("A26:D26")
    ws["A26"] = (
        f"Tannarx qamrovi: {coverage:.1f}%. Manba: Uzum purchasePrice. Yetishmagan qiymatlar nol bilan almashtirilmaydi."
        if uz
        else f"Покрытие себестоимостью: {coverage:.1f}%. Источник: Uzum purchasePrice. Отсутствующие значения не заменяются нулём."
    )
    ws["A26"].font = Font(name="Arial", size=8.5, italic=True, color="6B7280")
    ws["A26"].alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[26].height = 31

    section_title(28, "Buyurtmalar va sxemalar" if uz else "Заказы и схемы")
    operational_rows = [
        ("Berilgan buyurtmalar" if uz else "Выдано заказов", int(issued.get("orders") or 0), "Bekor qilingan buyurtmalar" if uz else "Отменено заказов", int(payload.get("cancellations") or 0)),
        ("FBO", int(schemes.get("FBO") or 0), "FBS", int(schemes.get("FBS") or 0)),
        ("DBS", int(schemes.get("DBS") or 0), "Uzum sxemani bermadi" if uz else "Схема не передана", int(schemes.get("UNKNOWN") or 0)),
    ]
    for row, values in enumerate(operational_rows, start=29):
        for column, value in enumerate(values, start=1):
            ws.cell(row, column, value)
            ws.cell(row, column).border = Border(bottom=GRID)
            ws.cell(row, column).alignment = Alignment(vertical="center", wrap_text=True)
            ws.cell(row, column).font = Font(name="Arial", size=9, bold=column in (1, 3), color="566573" if column in (1, 3) else NAVY)
        ws.cell(row, 2).alignment = Alignment(horizontal="right", vertical="center")
        ws.cell(row, 4).alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 24

    ws.column_dimensions["A"].width = 37
    ws.column_dimensions["B"].width = 19
    ws.column_dimensions["C"].width = 29
    ws.column_dimensions["D"].width = 27
    ws.freeze_panes = "A4"
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.print_area = "A1:D31"
    ws.sheet_view.zoomScale = 90

    ws = wb.create_sheet("Tovarlar" if uz else "По товарам")
    ws.sheet_view.showGridLines = False
    headers = (
        ["Tovar nomi", "Sxema", "Soni", "Daromad", "Chiqarishga", "Tannarx", "Soliq", "Tovar foydasi", "ROI %"]
        if uz
        else ["Товар", "Схема", "Кол-во", "Выручка", "К выплате", "Себестоимость", "Налог", "Прибыль товара", "ROI %"]
    )

    ws.merge_cells("A1:I1")
    ws["A1"] = "Tovarlar bo‘yicha natija" if uz else "Результат по товарам"
    ws["A1"].font = Font(name="Arial", size=16, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor=NAVY)
    ws["A1"].alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 32
    ws.merge_cells("A2:I2")
    ws["A2"] = (
        f"Sana: {report_date} · Do‘kon: {shop_label}"
        if uz
        else f"Дата: {report_date} · Магазин: {shop_label}"
    )
    ws["A2"].font = Font(name="Arial", size=10, color="566573")

    for column, header in enumerate(headers, start=1):
        cell = ws.cell(4, column, header)
        cell.font = Font(name="Arial", size=9, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=GRID)
    ws.row_dimensions[4].height = 30

    first_data_row = 5
    for row_index, product in enumerate(products, start=first_data_row):
        has_cost = product.get("cost_per_unit") is not None and float(product.get("known_cost_qty") or product.get("qty") or 0) > 0
        sku_text = _identifier(product.get("sku"))
        values = [
            _clean(product.get("title") or sku_text or "—")
            + (f" - {sku_text}" if sku_text else ""),
            _clean(product.get("scheme") or "—"),
            _num(product.get("qty")) or 0,
            _num(product.get("revenue")) or 0,
            _num(product.get("payout")) or 0,
            _num(product.get("cost_total")) if has_cost else None,
            _num(product.get("tax_expense")) or 0,
            None,
            None,
        ]
        for column, value in enumerate(values, start=1):
            cell = ws.cell(row_index, column, value)
            cell.font = Font(name="Arial", size=8.5)
            cell.fill = PatternFill("solid", fgColor=WHITE if row_index % 2 == 0 else PALE_BLUE)
            cell.alignment = Alignment(
                horizontal="left" if column == 1 else "center",
                vertical="center",
                wrap_text=True,
            )
            cell.border = Border(bottom=GRID)
        if values[5] is not None:
            ws.cell(row_index, 8, f"=E{row_index}-F{row_index}-G{row_index}")
            ws.cell(row_index, 9, f'=IFERROR(H{row_index}/F{row_index},"")')
        ws.row_dimensions[row_index].height = 31

    if products:
        last_data_row = first_data_row + len(products) - 1
        total_row = last_data_row + 1
        ws.cell(total_row, 1, "Umumiy:" if uz else "Итого:")
        for column in range(3, 9):
            letter = get_column_letter(column)
            ws.cell(total_row, column, f"=SUM({letter}{first_data_row}:{letter}{last_data_row})")
        ws.cell(total_row, 9, f'=IFERROR(H{total_row}/F{total_row},"")')
        for cell in ws[total_row]:
            cell.font = Font(name="Arial", size=9, bold=True)
            cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(top=GRID, bottom=GRID)
        note_row = total_row + 2
        ws.auto_filter.ref = f"A4:I{last_data_row}"
        ws.conditional_formatting.add(
            f"H{first_data_row}:H{last_data_row}",
            CellIsRule(operator="lessThan", formula=["0"], font=Font(color=RED), fill=PatternFill("solid", fgColor="FCE4D6")),
        )
        ws.conditional_formatting.add(
            f"I{first_data_row}:I{last_data_row}",
            CellIsRule(operator="lessThan", formula=["0"], font=Font(color=RED), fill=PatternFill("solid", fgColor="FCE4D6")),
        )
    else:
        last_data_row = 4
        total_row = 4
        note_row = 8
        ws.merge_cells("A5:I7")
        ws["A5"] = (
            "Tanlangan sana uchun savdo topilmadi. Agar savdo Uzum kabinetida endi paydo bo‘lgan bo‘lsa, u hisobotda biroz keyin ko‘rinishi mumkin."
            if uz
            else "За выбранную дату продажи не найдены. Если продажа только появилась в кабинете Uzum, она может отобразиться в отчёте немного позже."
        )
        ws["A5"].font = Font(name="Arial", size=11, bold=True, color="566573")
        ws["A5"].fill = PatternFill("solid", fgColor=PALE_BLUE)
        ws["A5"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=9)
    ws.cell(note_row, 1, (
        "Tovar foydasi faqat to‘lov, tannarx va soliqni hisobga oladi. Uzum qo‘shimcha xarajatlari va qaytarishlar «Xulosa» varag‘idagi yakuniy natijada ko‘rsatilgan."
        if uz
        else "Прибыль товара учитывает выплату, себестоимость и налог. Дополнительные расходы и возвраты Uzum входят в итог на листе «Сводка»."
    ))
    ws.cell(note_row, 1).font = Font(name="Arial", size=8.5, italic=True, color="6B7280")
    ws.cell(note_row, 1).alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[note_row].height = 31
    for column in ("D", "E", "F", "G", "H"):
        for cell in ws[column][first_data_row - 1 : max(note_row, total_row)]:
            cell.number_format = money_format
    for cell in ws["I"][first_data_row - 1 : total_row]:
        cell.number_format = percent_format
    _style_report_sheet(ws, widths=[46, 11, 9, 16, 16, 16, 14, 18, 12], freeze="A5")
    ws.print_title_rows = "4:4"
    ws.sheet_view.zoomScale = 85

    ws = wb.create_sheet("Uzum xarajatlari" if uz else "Расходы Uzum")
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:C1")
    ws["A1"] = "Uzum xarajatlari tafsiloti" if uz else "Расшифровка расходов Uzum"
    ws["A1"].font = Font(name="Arial", size=16, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor=NAVY)
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 32
    ws.merge_cells("A2:C2")
    ws["A2"] = f"Sana: {report_date}" if uz else f"Дата: {report_date}"
    ws["A2"].font = Font(name="Arial", size=10, color="566573")
    expense_headers = ["Operatsiya", "Yechim", "Qaytarish"] if uz else ["Операция", "Списание", "Возврат"]
    for column, value in enumerate(expense_headers, start=1):
        ws.cell(4, column, value)
        ws.cell(4, column).font = Font(name="Arial", size=9, bold=True, color=WHITE)
        ws.cell(4, column).fill = PatternFill("solid", fgColor=BLUE)
        ws.cell(4, column).alignment = Alignment(horizontal="center", vertical="center")
    if expenses:
        for row, (name, raw_amount) in enumerate(expenses, start=5):
            amount = float(raw_amount or 0)
            ws.cell(row, 1, _clean(name) or "—")
            ws.cell(row, 2, amount if amount > 0 else None)
            ws.cell(row, 3, abs(amount) if amount < 0 else None)
            for column in range(1, 4):
                ws.cell(row, column).border = Border(bottom=GRID)
                ws.cell(row, column).fill = PatternFill("solid", fgColor=WHITE if row % 2 == 0 else PALE_BLUE)
                ws.cell(row, column).alignment = Alignment(horizontal="left" if column == 1 else "right", vertical="center", wrap_text=True)
            ws.row_dimensions[row].height = 26
        expense_last_row = 4 + len(expenses)
        total_row = expense_last_row + 1
        ws.cell(total_row, 1, "Umumiy" if uz else "Итого")
        ws.cell(total_row, 2, f"=SUM(B5:B{expense_last_row})")
        ws.cell(total_row, 3, f"=SUM(C5:C{expense_last_row})")
        for cell in ws[total_row]:
            cell.font = Font(name="Arial", size=9, bold=True)
            cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
            cell.border = Border(top=GRID, bottom=GRID)
        ws.auto_filter.ref = f"A4:C{expense_last_row}"
        note_row = total_row + 2
    else:
        ws.merge_cells("A5:C7")
        ws["A5"] = (
            "Tanlangan sana uchun Uzum xarajatlari topilmadi."
            if uz
            else "За выбранную дату отдельные расходы Uzum не найдены."
        )
        ws["A5"].font = Font(name="Arial", size=11, bold=True, color="566573")
        ws["A5"].fill = PatternFill("solid", fgColor=PALE_BLUE)
        ws["A5"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        note_row = 8
    ws.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=3)
    ws.cell(note_row, 1, (
        "Komissiya va logistika bu ro‘yxatda takrorlanmaydi: ular «To‘lovga» summasida allaqachon hisobga olingan."
        if uz
        else "Комиссия и логистика здесь повторно не вычитаются: они уже учтены в сумме «К выплате»."
    ))
    ws.cell(note_row, 1).font = Font(name="Arial", size=8.5, italic=True, color="6B7280")
    ws.cell(note_row, 1).alignment = Alignment(wrap_text=True, vertical="center")
    ws.row_dimensions[note_row].height = 31
    for column in ("B", "C"):
        for cell in ws[column][4:note_row]:
            cell.number_format = money_format
    _style_report_sheet(ws, widths=[62, 20, 20], freeze="A5")
    ws.print_title_rows = "4:4"
    ws.sheet_view.zoomScale = 90

    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.save(path)
    wb.close()
    return path


def _pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    app_dir = Path(__file__).resolve().parent
    regular_candidates = [
        app_dir / "DejaVuSans.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
    ]
    bold_candidates = [
        app_dir / "DejaVuSans-Bold.ttf",
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf"),
    ]
    regular_path = next((item for item in regular_candidates if item.exists()), None)
    bold_path = next((item for item in bold_candidates if item.exists()), regular_path)
    if regular_path is None or bold_path is None:
        raise RuntimeError("Unicode PDF font was not found")
    if "SellerProDaily" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("SellerProDaily", str(regular_path)))
    if "SellerProDailyBold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("SellerProDailyBold", str(bold_path)))
    return "SellerProDaily", "SellerProDailyBold"


def build_daily_finance_pdf(
    payload: dict[str, Any],
    output: str | Path | None = None,
    *,
    lang: str = "uz",
) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    path = _resolve_output(output, "sellerpro_daily_", ".pdf")
    regular, bold = _pdf_fonts()
    uz = str(lang).lower() == "uz"
    products = list(payload.get("products") or [])
    report_date = _date_label(payload.get("date"))
    shop_label = _clean(payload.get("shop_label") or payload.get("shop_id") or "—")
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title=f"SellerPro daily report {report_date}",
    )
    title_style = ParagraphStyle("DailyTitle", fontName=bold, fontSize=16, leading=19, textColor=colors.HexColor("#" + NAVY), alignment=TA_LEFT)
    section_style = ParagraphStyle("DailySection", fontName=bold, fontSize=10, leading=13, textColor=colors.HexColor("#" + NAVY), spaceAfter=4)
    meta_style = ParagraphStyle("DailyMeta", fontName=regular, fontSize=8, leading=10, textColor=colors.HexColor("#566573"))
    note_style = ParagraphStyle("DailyNote", parent=meta_style, fontSize=7.5, leading=9)
    cell_style = ParagraphStyle("DailyCell", fontName=regular, fontSize=6.4, leading=8)
    cell_center = ParagraphStyle("DailyCellCenter", parent=cell_style, alignment=1)
    summary_label = ParagraphStyle("DailySummaryLabel", fontName=regular, fontSize=7.5, leading=9)
    summary_bold = ParagraphStyle("DailySummaryBold", parent=summary_label, fontName=bold)
    summary_value = ParagraphStyle("DailySummaryValue", parent=summary_label, alignment=2)
    summary_value_bold = ParagraphStyle("DailySummaryValueBold", parent=summary_value, fontName=bold)
    header_style = ParagraphStyle("DailyHeader", parent=cell_center, fontName=bold, fontSize=6.4, textColor=colors.white)
    total_style = ParagraphStyle("DailyTotal", parent=cell_center, fontName=bold, fontSize=6.4, textColor=colors.HexColor("#333333"))
    bridge = _daily_bridge_values(payload, products)
    final = bridge["final"]
    complete = bool(payload.get("complete"))
    coverage = float(payload.get("cost_coverage") or 0) * 100
    accepted = dict(payload.get("accepted") or {})
    issued = dict(payload.get("issued") or {})
    schemes = dict(payload.get("schemes") or {})
    result_label = (
        "Yakuniy sof foyda"
        if uz and complete
        else "Ma’lum ma’lumotlar bo‘yicha natija"
        if uz
        else "Итоговая чистая прибыль"
        if complete
        else "Результат по известным данным"
    )
    story: list[Any] = [
        Paragraph(("Seller.pro.uz · Kunlik hisobot" if uz else "Seller.pro.uz · Дневной отчёт"), title_style),
        Paragraph(
            (
                f"Sana: {report_date} · Do‘kon: {xml_escape(shop_label)} · Valyuta: so‘m"
                if uz
                else f"Дата: {report_date} · Магазин: {xml_escape(shop_label)} · Валюта: сум"
            ),
            meta_style,
        ),
        Spacer(1, 4 * mm),
        Paragraph("Kun yakuni" if uz else "Главное за день", section_style),
    ]

    roi_text = "—" if bridge["cost"] <= 0 else f"{final / bridge['cost'] * 100:.1f}%"
    kpi_data = [
        [
            Paragraph("Qabul qilingan buyurtmalar" if uz else "Принято заказов", summary_label),
            Paragraph(str(int(accepted.get("orders") or 0)), summary_value_bold),
            Paragraph("Sotilgan tovarlar" if uz else "Продано товаров", summary_label),
            Paragraph(f"{float(accepted.get('units') or 0):g}", summary_value_bold),
        ],
        [
            Paragraph("Tushum" if uz else "Выручка", summary_label),
            Paragraph(_money(bridge["revenue"]), summary_value_bold),
            Paragraph("To‘lovga" if uz else "К выплате", summary_label),
            Paragraph(_money(bridge["payout"]), summary_value_bold),
        ],
        [
            Paragraph(result_label, summary_bold),
            Paragraph(_money(final), summary_value_bold),
            Paragraph("ROI", summary_bold),
            Paragraph(roi_text, summary_value_bold),
        ],
    ]
    kpi_table = Table(kpi_data, colWidths=[80 * mm, 55 * mm, 80 * mm, 55 * mm], hAlign="LEFT")
    kpi_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 1), colors.HexColor("#" + PALE_BLUE)),
                ("BACKGROUND", (0, 2), (-1, 2), colors.HexColor("#E2F0D9" if final >= 0 else "#FCE4D6")),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E2EA")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EA")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([kpi_table, Spacer(1, 4 * mm), Paragraph("Natija qanday hisoblandi" if uz else "Как получился итог", section_style)])

    bridge_rows = [
        ("1. Tushum" if uz else "1. Выручка", _money(bridge["revenue"]), "Sotilgan tovarlar narxi" if uz else "Цена проданных товаров", False, "normal"),
        ("2. − Uzum komissiyasi" if uz else "2. − Комиссия Uzum", _signed_money(bridge["commission"], "-"), "To‘lovda hisobga olingan" if uz else "Уже учтена в сумме «К выплате»", False, "negative"),
        ("3. − Logistika" if uz else "3. − Логистика", _signed_money(bridge["logistics"], "-"), "To‘lovda hisobga olingan" if uz else "Уже учтена в сумме «К выплате»", False, "negative"),
        ("4. − Boshqa ushlanmalar" if uz else "4. − Другие удержания", _signed_money(bridge["other_payout"], "-"), "To‘lov ichidagi boshqa ushlanmalar" if uz else "Другие удержания внутри выплаты", False, "negative"),
        ("5. + To‘lov tuzatishi" if uz else "5. + Корректировка выплаты", _signed_money(bridge["payout_adjustment"], "+"), "Uzum bergan to‘lov bilan farq" if uz else "Разница с фактической выплатой Uzum", False, "positive"),
        ("= To‘lovga" if uz else "= К выплате", _money(bridge["payout"]), "Uzum ushlanmalaridan keyin" if uz else "После удержаний Uzum", True, "normal"),
        ("6. − Tannarx" if uz else "6. − Себестоимость", _signed_money(bridge["cost"], "-"), "Uzum purchasePrice qiymati" if uz else "Закупочная стоимость из Uzum purchasePrice", False, "negative"),
        ("= Soliqdan oldingi foyda" if uz else "= Прибыль до налога", _money(bridge["profit_before_tax"]), "To‘lov minus tannarx" if uz else "К выплате минус себестоимость", True, "normal"),
        ("7. − Soliq" if uz else "7. − Налог", _signed_money(bridge["tax"], "-"), "Bot sozlamalaridagi stavka" if uz else "По ставке из настроек бота", False, "negative"),
        ("= Tovarlar foydasi" if uz else "= Прибыль товаров", _money(bridge["product_profit"]), "Qo‘shimcha xarajatlardan oldin" if uz else "До дополнительных расходов", True, "normal"),
        ("8. − Uzum qo‘shimcha xarajatlari" if uz else "8. − Доп. расходы Uzum", _signed_money(bridge["deductions"], "-"), "Saqlash, reklama va boshqa xizmatlar" if uz else "Хранение, продвижение и другие услуги", False, "negative"),
        ("9. + Uzum qaytargan mablag‘" if uz else "9. + Возвраты Uzum", _signed_money(bridge["refunds"], "+"), "Uzum qaytargan summa" if uz else "Средства, возвращённые Uzum", False, "positive"),
        ("10. − Tashqi xarajatlar" if uz else "10. − Внешние расходы", _signed_money(bridge["external"], "-"), "Uzumdan tashqaridagi xarajatlar" if uz else "Расходы вне Uzum, добавленные вручную", False, "negative"),
        (f"= {result_label}", _money(final), "Barcha ma’lum xarajatlardan keyin" if uz else "После всех известных расходов", True, "final"),
    ]
    summary_data: list[list[Any]] = [
        [
            Paragraph("Qadam" if uz else "Шаг", summary_bold),
            Paragraph("Summa" if uz else "Сумма", summary_value_bold),
            Paragraph("Izoh" if uz else "Что означает", summary_bold),
        ]
    ]
    for label, amount, explanation, subtotal, _ in bridge_rows:
        summary_data.append(
            [
                Paragraph(label, summary_bold if subtotal else summary_label),
                Paragraph(amount, summary_value_bold if subtotal else summary_value),
                Paragraph(explanation, summary_label),
            ]
        )
    summary_table = Table(summary_data, colWidths=[90 * mm, 40 * mm, 140 * mm], repeatRows=1, hAlign="LEFT")
    summary_styles: list[tuple[Any, ...]] = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLUE)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EA")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for index, (_, _, _, subtotal, kind) in enumerate(bridge_rows, start=1):
        if subtotal:
            summary_styles.append(("BACKGROUND", (0, index), (-1, index), colors.HexColor("#" + PALE_BLUE)))
            summary_styles.append(("LINEABOVE", (0, index), (-1, index), 0.5, colors.HexColor("#" + BLUE)))
        if kind == "negative":
            summary_styles.append(("TEXTCOLOR", (1, index), (1, index), colors.HexColor("#" + RED)))
        elif kind == "positive":
            summary_styles.append(("TEXTCOLOR", (1, index), (1, index), colors.HexColor("#" + GREEN)))
        elif kind == "final":
            summary_styles.append(("BACKGROUND", (0, index), (-1, index), colors.HexColor("#E2F0D9" if final >= 0 else "#FCE4D6")))
            summary_styles.append(("LINEABOVE", (0, index), (-1, index), 1, colors.HexColor("#" + BLUE)))
            summary_styles.append(("TEXTCOLOR", (1, index), (1, index), colors.HexColor("#" + (GREEN if final >= 0 else RED))))
    summary_table.setStyle(TableStyle(summary_styles))
    note = (
        f"Tannarx qamrovi: {coverage:.1f}%. Manba: Uzum purchasePrice. Yetishmagan qiymatlar nol bilan almashtirilmaydi."
        if uz
        else f"Покрытие себестоимостью: {coverage:.1f}%. Источник: Uzum purchasePrice. Отсутствующие значения не заменяются нулём."
    )
    story.extend([summary_table, Spacer(1, 2 * mm), Paragraph(note, note_style)])

    operational_data = [
        [
            Paragraph("Berilgan buyurtmalar" if uz else "Выдано заказов", summary_label),
            Paragraph(str(int(issued.get("orders") or 0)), summary_value_bold),
            Paragraph("Bekor qilingan" if uz else "Отменено заказов", summary_label),
            Paragraph(str(int(payload.get("cancellations") or 0)), summary_value_bold),
        ],
        [
            Paragraph("Sxemalar" if uz else "Схемы", summary_label),
            Paragraph(f"FBO {int(schemes.get('FBO') or 0)} · FBS {int(schemes.get('FBS') or 0)}", summary_value),
            Paragraph("DBS / noma’lum" if uz else "DBS / не передано", summary_label),
            Paragraph(f"DBS {int(schemes.get('DBS') or 0)} · {int(schemes.get('UNKNOWN') or 0)}", summary_value),
        ],
    ]
    operational_table = Table(operational_data, colWidths=[70 * mm, 65 * mm, 70 * mm, 65 * mm], hAlign="LEFT")
    operational_table.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.3, colors.HexColor("#D9E2EA")), ("INNERGRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#D9E2EA")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
    story.extend([Spacer(1, 3 * mm), operational_table])

    if not products:
        no_data = Table(
            [[Paragraph(
                "Tanlangan sana uchun savdo topilmadi. Agar savdo Uzum kabinetida endi paydo bo‘lgan bo‘lsa, u hisobotda biroz keyin ko‘rinishi mumkin."
                if uz
                else "За выбранную дату продажи не найдены. Если продажа только появилась в кабинете Uzum, она может отобразиться в отчёте немного позже.",
                summary_bold,
            )]],
            colWidths=[270 * mm],
            hAlign="LEFT",
        )
        no_data.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#" + PALE_BLUE)), ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E2EA")), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
        story.extend([Spacer(1, 4 * mm), no_data])
        doc.build(story)
        return path

    story.extend([PageBreak(), Paragraph("Tovarlar bo‘yicha natija" if uz else "Результат по товарам", section_style)])
    headers = (
        ["Tovar nomi", "Sxema", "Soni", "Daromad", "Chiqarishga", "Tannarx", "Soliq", "Tovar foydasi", "ROI %"]
        if uz
        else ["Товар", "Схема", "Кол-во", "Выручка", "К выплате", "Себестоимость", "Налог", "Прибыль товара", "ROI %"]
    )
    header_row: list[Any] = [Paragraph(value, header_style) for value in headers]
    product_rows: list[list[Any]] = []
    for product in products:
        title = _clean(product.get("title") or "—")
        sku = _identifier(product.get("sku"))
        if sku:
            title += f" - {sku}"
        has_cost = product.get("cost_per_unit") is not None and float(product.get("known_cost_qty") or product.get("qty") or 0) > 0
        cost = _num(product.get("cost_total")) if has_cost else None
        tax = _num(product.get("tax_expense")) or 0.0
        payout = _num(product.get("payout")) or 0.0
        net = payout - cost - tax if cost is not None else None
        roi = net / cost * 100.0 if net is not None and cost and cost > 0 else None
        product_rows.append(
            [
                Paragraph(xml_escape(title), cell_style),
                Paragraph(xml_escape(_clean(product.get("scheme") or "—")), cell_center),
                Paragraph(f"{float(product.get('qty') or 0):g}", cell_center),
                Paragraph(_money(product.get("revenue")), cell_center),
                Paragraph(_money(payout), cell_center),
                Paragraph(_money(cost), cell_center),
                Paragraph(_money(tax), cell_center),
                Paragraph(_money(net), cell_center),
                Paragraph("—" if roi is None else f"{roi:.2f}%", cell_center),
            ]
        )
    totals = {
        key: sum(float(item.get(key) or 0) for item in products)
        for key in ("qty", "revenue", "payout", "cost_total", "tax_expense")
    }
    known_products = [
        item
        for item in products
        if item.get("cost_per_unit") is not None
        and float(item.get("known_cost_qty") or item.get("qty") or 0) > 0
    ]
    known_revenue = sum(
        float(item.get("known_revenue") if item.get("known_revenue") is not None else item.get("revenue") or 0)
        for item in known_products
    )
    known_payout = sum(
        float(item.get("known_payout") if item.get("known_payout") is not None else item.get("payout") or 0)
        for item in known_products
    )
    known_tax = sum(
        float(item.get("known_tax_expense") if item.get("known_tax_expense") is not None else item.get("tax_expense") or 0)
        for item in known_products
    )
    known_net = sum(
        float(item.get("known_payout") if item.get("known_payout") is not None else item.get("payout") or 0)
        - float(item.get("cost_total") or 0)
        - float(item.get("known_tax_expense") if item.get("known_tax_expense") is not None else item.get("tax_expense") or 0)
        for item in known_products
    )
    roi_total = known_net / totals["cost_total"] * 100.0 if totals["cost_total"] > 0 else None
    total_pdf_row = [
        Paragraph("Umumiy:" if uz else "Итого:", total_style),
        "",
        Paragraph(f"{totals['qty']:g}", total_style),
        Paragraph(_money(totals["revenue"]), total_style),
        Paragraph(_money(totals["payout"]), total_style),
        Paragraph(_money(totals["cost_total"]), total_style),
        Paragraph(_money(totals["tax_expense"]), total_style),
        Paragraph(_money(known_net), total_style),
        Paragraph("—" if roi_total is None else f"{roi_total:.2f}%", total_style),
    ]
    # Fixed-size chunks avoid leaving one orphan product row next to the final
    # summary.  Each chunk is a separate table, so its header is guaranteed.
    page_size = 21
    chunks = [product_rows[index : index + page_size] for index in range(0, len(product_rows), page_size)] or [[]]
    widths = [103 * mm, 15 * mm, 12 * mm, 25 * mm, 25 * mm, 25 * mm, 21 * mm, 26 * mm, 18 * mm]
    for chunk_index, chunk in enumerate(chunks):
        last_chunk = chunk_index == len(chunks) - 1
        data = [header_row, *chunk]
        if last_chunk:
            data.append(total_pdf_row)
        table = Table(data, colWidths=widths, repeatRows=1)
        body_end = -2 if last_chunk else -1
        styles = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BLUE)),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9E2EA")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        if chunk:
            styles.append(("ROWBACKGROUNDS", (0, 1), (-1, body_end), [colors.white, colors.HexColor("#" + PALE_BLUE)]))
        if last_chunk:
            styles.append(("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#" + PALE_GRAY)))
        table.setStyle(TableStyle(styles))
        story.append(table)
        if last_chunk:
            story.append(Spacer(1, 4 * mm))
        else:
            story.append(PageBreak())
    story.extend([
        Spacer(1, 2 * mm),
        Paragraph(
            "Tovar foydasi to‘lov, tannarx va soliqni hisobga oladi. Uzum qo‘shimcha xarajatlari va qaytarishlar birinchi sahifadagi yakuniy natijada ko‘rsatilgan."
            if uz
            else "Прибыль товара учитывает выплату, себестоимость и налог. Дополнительные расходы и возвраты Uzum входят в итог на первой странице.",
            note_style,
        ),
    ])
    doc.build(story)
    return path


def build_compensation_workbook(
    rows: list[dict[str, Any]],
    output: str | Path | None = None,
    *,
    kind: str,
    lang: str = "ru",
) -> Path:
    normalized_kind = str(kind).lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    path = _resolve_output(output, f"sellerpro_{normalized_kind}_", ".xlsx")
    uz = str(lang).lower() == "uz"
    wb = Workbook()
    ws = wb.active
    ws.title = (
        "Yo‘qolgan tovarlar" if uz and normalized_kind == "loss" else
        "Brak hisoboti" if uz else
        "Потерянные товары" if normalized_kind == "loss" else
        "Отчет браков"
    )
    if normalized_kind == "loss":
        headers = (
            ["Tovar nomi", "Shtrix-kod", "Qoplama miqdori", "Soni", "Jami"]
            if uz
            else ["Название товара", "Штрих-код", "Размер возмещения", "Количество", "Итого"]
        )
        widths = [82, 22, 22, 14, 22]
    else:
        headers = (
            ["Tovar ID", "Tovar nomi", "SKU", "Shtrix-kod", "Qoplama miqdori", "Soni", "Jami"]
            if uz
            else ["ИД товара", "Название товара", "SKU", "ШК", "Размер возмещения", "Кол-во", "Итого"]
        )
        widths = [16, 65, 24, 22, 22, 12, 22]
    total_col = len(headers)
    total_qty = sum(_qty(item.get("quantity")) for item in rows)
    known_total = sum(float(_num(item.get("total")) or 0) for item in rows)
    incomplete = sum(
        1
        for item in rows
        if _num(item.get("compensation")) is None or _num(item.get("total")) is None
    )
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_col)
    ws.cell(1, 1, "Yo‘qolgan tovarlar reestri" if uz and normalized_kind == "loss" else "Brak tovarlar reestri" if uz else "Реестр потерянных товаров" if normalized_kind == "loss" else "Реестр повреждённых товаров")
    ws.cell(1, 1).font = Font(name="Arial", size=16, bold=True, color=NAVY)
    ws.cell(1, 1).alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 30
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=total_col)
    ws.cell(2, 1, (
        f"Pozitsiyalar: {len(rows)} · Jami birlik: {total_qty} · Hisoblangan summa: {_money(known_total)}"
        if uz
        else f"Позиций: {len(rows)} · Единиц: {total_qty} · Рассчитанная сумма: {_money(known_total)}"
    ))
    ws.cell(2, 1).font = Font(name="Arial", size=10, bold=True, color="34495E")
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=total_col)
    ws.cell(3, 1, (
        f"Diqqat: {incomplete} qator bo‘yicha aniq narx yoki komissiya yo‘q; summa bo‘sh qoldirildi."
        if uz and incomplete
        else "Ma’lumotlar Uzum API asosida. Summalar taxmin qilinmaydi."
        if uz
        else f"Внимание: по {incomplete} строкам нет точной цены или комиссии; сумма оставлена пустой."
        if incomplete
        else "Данные получены из Uzum API. Отсутствующие суммы не подставляются приблизительно."
    ))
    ws.cell(3, 1).font = Font(name="Arial", size=9, italic=True, color=RED if incomplete else "6B7280")
    header_row = 4
    for column, header in enumerate(headers, start=1):
        ws.cell(header_row, column, header)
    for cell in ws[header_row]:
        cell.font = Font(name="Arial", bold=True, color="333333")
        cell.fill = PatternFill("solid", fgColor="E7E7E7")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=GRID)
    ws.row_dimensions[header_row].height = 28

    for index, item in enumerate(rows, start=header_row + 1):
        compensation = _num(item.get("compensation"))
        if normalized_kind == "loss":
            values = [item.get("title"), item.get("barcode"), compensation, item.get("quantity"), item.get("total")]
        else:
            values = [item.get("product_id"), item.get("title"), item.get("sku_id") or item.get("sku"), item.get("barcode"), compensation, item.get("quantity"), item.get("total")]
        ws.append(values)
        for cell in ws[index]:
            cell.fill = PatternFill("solid", fgColor=WHITE if index % 2 == 0 else PALE_BLUE)
            cell.font = Font(name="Arial", size=9)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(bottom=GRID)
        ws.row_dimensions[index].height = 34

    total_row = header_row + len(rows) + 1
    qty_col = total_col - 1
    ws.cell(total_row, 1, "Umumiy:" if uz else "Итого:")
    ws.cell(total_row, qty_col, f"=SUM({get_column_letter(qty_col)}{header_row + 1}:{get_column_letter(qty_col)}{total_row - 1})")
    ws.cell(total_row, total_col, f"=SUM({get_column_letter(total_col)}{header_row + 1}:{get_column_letter(total_col)}{total_row - 1})")
    for cell in ws[total_row]:
        cell.font = Font(name="Arial", bold=True)
        cell.fill = PatternFill("solid", fgColor=PALE_GRAY)
        cell.border = Border(top=GRID, bottom=GRID)
    for column in (total_col - 2, total_col):
        for cell in ws[get_column_letter(column)][header_row:]:
            cell.number_format = '#,##0'
    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(total_col)}{max(header_row, total_row - 1)}"
    ws.sheet_view.showGridLines = False
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.page_setup.orientation = "landscape"
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.print_title_rows = f"{header_row}:{header_row}"
    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.save(path)
    wb.close()
    return path


def _set_docx_run_font(run: Any, *, bold: bool | None = None, size: float | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def build_claim_docx(
    rows: list[dict[str, Any]],
    output: str | Path | None = None,
    *,
    kind: str,
    shop_id: int | str | None = None,
    seller_details: dict[str, Any] | None = None,
    lang: str = "ru",
) -> Path:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    normalized_kind = str(kind).lower()
    if normalized_kind not in {"loss", "damage"}:
        raise ValueError("kind must be 'loss' or 'damage'")
    path = _resolve_output(output, f"sellerpro_claim_{normalized_kind}_", ".docx")
    uz = str(lang).lower() == "uz"
    details = dict(seller_details or {})
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark = header.add_run("⬡ ")
    _set_docx_run_font(mark, bold=True, size=15)
    mark.font.color.rgb = RGBColor(108, 49, 217)
    brand = header.add_run("UZUM MARKET")
    _set_docx_run_font(brand, bold=True, size=13)
    brand.font.color.rgb = RGBColor(108, 49, 217)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("PRETENZIYA" if uz else "ПРЕТЕНЗИЯ")
    _set_docx_run_font(run, bold=True, size=14)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(5)
    intro_text = (
        "Quyidagi tovarlar bo‘yicha kompensatsiya to‘lashni so‘rayman. "
        "Maydonlarni yuborishdan oldin sotuvchi ma’lumotlari bilan to‘ldiring."
        if uz
        else "Прошу выплатить компенсацию по перечисленным ниже товарам. "
        "Перед отправкой заполните реквизиты продавца."
    )
    run = intro.add_run(intro_text)
    _set_docx_run_font(run, size=8.5)

    labels = (
        [
            ("Sotuvchi F.I.Sh.", "full_name"),
            ("YaTT / tashkilot", "company"),
            ("JShShIR / STIR", "tax_id"),
            ("Hisob raqami", "bank_account"),
            ("MFO", "mfo"),
            ("Do‘kon ID", "shop_id"),
        ]
        if uz
        else [
            ("Ф.И.О. продавца", "full_name"),
            ("ИП / организация", "company"),
            ("ПИНФЛ / ИНН", "tax_id"),
            ("Расчётный счёт", "bank_account"),
            ("МФО", "mfo"),
            ("ID магазина", "shop_id"),
        ]
    )
    details["shop_id"] = details.get("shop_id") or shop_id or ""
    details_table = document.add_table(rows=3, cols=4)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details_table.autofit = False
    widths = [Cm(3.2), Cm(6.2), Cm(3.0), Cm(6.0)]
    for index, (label, key) in enumerate(labels):
        table_row = index // 2
        pair = (index % 2) * 2
        details_table.cell(table_row, pair).width = widths[pair]
        details_table.cell(table_row, pair + 1).width = widths[pair + 1]
        details_table.cell(table_row, pair).text = label
        details_table.cell(table_row, pair + 1).text = _clean(details.get(key)) or "____________________________"
    for row in details_table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for item in paragraph.runs:
                    _set_docx_run_font(item, bold=index % 2 == 0, size=7.5)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = "w:" + edge
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "4")
                element.set(qn("w:color"), "D9E2EA")

    preview_qty = sum(_qty(item.get("quantity")) for item in rows)
    preview_total = sum(float(_num(item.get("total")) or 0) for item in rows)
    preview_incomplete = sum(
        1
        for item in rows
        if _num(item.get("compensation")) is None or _num(item.get("total")) is None
    )
    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(5)
    summary.paragraph_format.space_after = Pt(5)
    summary_text = (
        f"Hujjat tarkibi: {len(rows)} pozitsiya · {preview_qty} dona · hisoblangan summa {_money(preview_total)} so‘m."
        if uz
        else f"Состав претензии: {len(rows)} позиций · {preview_qty} единиц · рассчитанная сумма {_money(preview_total)} сум."
    )
    if preview_incomplete:
        summary_text += (
            f" {preview_incomplete} qator bo‘yicha summa ma’lumot yetishmagani sababli bo‘sh qoldirilgan."
            if uz
            else f" По {preview_incomplete} строкам сумма оставлена пустой из-за отсутствующих данных."
        )
    summary_run = summary.add_run(summary_text)
    _set_docx_run_font(summary_run, bold=True, size=8.2)
    summary_run.font.color.rgb = RGBColor(23, 54, 93)

    reason = "Yo‘qotish" if uz and normalized_kind == "loss" else "Shikastlanish" if uz else "Утеря" if normalized_kind == "loss" else "Повреждение"
    headers = (
        ["Tovar nomi", "Shtrix-kod", "Sabab", "Qoplama*", "Soni", "Jami"]
        if uz
        else ["Название товара", "Штрих-код", "Причина", "Размер возмещения*", "Кол-во", "Итого"]
    )
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_widths = [Cm(8.0), Cm(2.7), Cm(2.2), Cm(2.6), Cm(1.2), Cm(2.5)]
    for index, header_text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = column_widths[index]
        cell.text = header_text
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E7E7E7")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item in paragraph.runs:
                _set_docx_run_font(item, bold=True, size=6.6)

    known_total = 0.0
    total_qty = 0
    incomplete = 0
    for index, item in enumerate(rows, start=1):
        row = table.add_row()
        compensation = _num(item.get("compensation"))
        total = _num(item.get("total"))
        if compensation is None or total is None:
            incomplete += 1
        else:
            known_total += total
        quantity = _qty(item.get("quantity"))
        total_qty += quantity
        title_text = _clean(item.get("title"))
        sku_text = _identifier(item.get("sku"))
        if sku_text:
            title_text += f" - {sku_text}"
        values = [
            title_text,
            _identifier(item.get("barcode")) or "—",
            reason,
            _money(compensation),
            str(quantity),
            _money(total),
        ]
        for column, value in enumerate(values):
            cell = row.cells[column]
            cell.width = column_widths[column]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if index % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), PALE_BLUE)
                cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for item_run in paragraph.runs:
                    _set_docx_run_font(item_run, size=6.2)
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)

    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[3])
    total_row.cells[0].text = "Umumiy:" if uz else "Итого:"
    total_row.cells[4].text = str(total_qty)
    total_row.cells[5].text = _money(known_total)
    for cell in total_row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E7E7E7")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item in paragraph.runs:
                _set_docx_run_font(item, bold=True, size=7)

    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().insert(0, table_header)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(2)
    note_text = (
        "* Qoplama miqdori Uzumdagi joriy sotuv narxidan SKU komissiyasi ayrilgan holda hisoblandi."
        if uz
        else "* Размер возмещения рассчитан как текущая стоимость товара в Uzum за вычетом комиссии маркетплейса по SKU."
    )
    if incomplete:
        note_text += (
            f" {incomplete} ta SKU bo‘yicha narx yoki aniq komissiya yo‘q; qiymat taxmin qilinmadi."
            if uz
            else f" По {incomplete} SKU нет цены или точной комиссии; сумма не рассчитана и оставлена пустой."
        )
    run = note.add_run(note_text)
    _set_docx_run_font(run, size=7.3)
    run.italic = True

    signature = document.add_paragraph()
    signature.paragraph_format.space_before = Pt(8)
    signature.add_run(("Sana: ______________    Imzo: ____________________" if uz else "Дата: ______________    Подпись: ____________________"))
    for item in signature.runs:
        _set_docx_run_font(item, size=8.5)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run("Seller.pro.uz · Uzum API ma’lumotlari" if uz else "Seller.pro.uz · данные Uzum API")
    _set_docx_run_font(footer_run, size=7)
    footer_run.font.color.rgb = RGBColor(128, 139, 150)

    document.save(path)
    return path
